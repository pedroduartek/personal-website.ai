# -*- coding: utf-8 -*-
"""Builds Pedro's CV (.docx + .pdf) from scripts/cv/cv-build.json.

Reproduces the current CV's structure and style (Word / Aptos): name 36pt bold,
16pt bold section headers, 12pt body, teal links, company/institution logos in
the left column, `**bold**` emphasis inside prose, two-column entries for
Experience/Education/Conferences. Run the exporter first (npm run cv does both).

Needs: python-docx, Pillow (webp->png for logos), pywin32 + Microsoft Word (PDF).
"""
import json
import os
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, "cv-build.json")
OUT_DIR = os.path.normpath(os.path.join(HERE, "..", "..", "src", "CV"))
DOCX = os.path.join(OUT_DIR, "Pedro_Duarte_CV.docx")
PDF = os.path.join(OUT_DIR, "Pedro_Duarte_CV.pdf")

FONT = "Aptos"
TEAL = RGBColor(0x46, 0x78, 0x86)
BODY = 12
HEADER = 16
NAME = 36
LEFT_COL = Inches(1.2)
RIGHT_COL = Inches(5.3)
LOGO_W = Inches(0.7)


def fmt_date(ym):
    if not ym:
        return "Present"
    try:
        return datetime.strptime(ym, "%Y-%m").strftime("%b %Y")
    except ValueError:
        return ym


def add_run(p, text, *, bold=False, italic=False, size=BODY, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def add_rich(p, text, *, size=BODY):
    """Render text with **bold** markers: even segments normal, odd segments bold."""
    for i, seg in enumerate(text.split("**")):
        if seg:
            add_run(p, seg, bold=(i % 2 == 1), size=size)


def logo_png(path):
    """Convert a source logo (webp) to a temp PNG python-docx can embed."""
    if not path or not os.path.exists(path):
        return None
    try:
        png = os.path.join(
            tempfile.gettempdir(),
            "_cvlogo_" + os.path.splitext(os.path.basename(path))[0] + ".png",
        )
        Image.open(path).convert("RGBA").save(png)
        return png
    except Exception as e:  # noqa: BLE001
        print("logo skipped:", path, e)
        return None


def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True  # header stays with its content
    add_run(p, text, bold=True, size=HEADER)


def two_col(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    # Fixed layout so Word honours the narrow left column (forces the 2-word
    # labels to wrap) instead of auto-sizing to content.
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    left, right = table.rows[0].cells
    left.width = LEFT_COL
    right.width = RIGHT_COL
    return left, right


def left_label(cell, name, logo=None, sub=None):
    """Left column: logo + name in ONE paragraph (so a page break can never
    separate them), then an optional sub-line."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    png = logo_png(logo)
    if png:
        p.add_run().add_picture(png, width=LOGO_W)
        p.add_run().add_break()  # line break inside the same (unbreakable) paragraph
    add_run(p, name, bold=True)
    if sub:
        sp = cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.keep_with_next = True
        add_run(sp, sub)


def rline(cell, first=False, space_after=2):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build():
    with open(DATA, encoding="utf-8") as f:
        cv = json.load(f)

    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(BODY)
    for s in doc.sections:
        s.page_width = Inches(8.5)
        s.page_height = Inches(11)
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)

    p = cv["profile"]

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    add_run(name_p, p["name"], bold=True, size=NAME)
    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(0)
    add_run(contact, f"{p['location']}  |  {p['email']}  |  {p['phone']}")
    link_p = doc.add_paragraph()
    add_run(link_p, p["linkedin"].replace("https://www.", "").rstrip("/"), color=TEAL)

    section_header(doc, "Summary")
    add_run(doc.add_paragraph(), cv["summary"])

    section_header(doc, "Technical Skills")
    for grp in cv["skills"]:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(3)
        add_run(bp, "• ")
        add_run(bp, f"{grp['group']}: ", bold=True)
        add_run(bp, grp["items"])

    section_header(doc, "Professional Experience")
    for e in cv["experience"]:
        left, right = two_col(doc)
        left_label(
            left,
            e["company"],
            logo=e.get("logo"),
            sub="(Current Company)" if e.get("current") else None,
        )
        add_run(rline(right, first=True), e["title"], bold=True)
        add_run(rline(right, space_after=6),
                f"{fmt_date(e['startDate'])} – {fmt_date(e['endDate'])}", italic=True)
        if e.get("technologies"):
            tp = rline(right, space_after=6)
            add_run(tp, "Tech stack: ", bold=True)
            add_run(tp, ", ".join(e["technologies"]))
        for para in e.get("prose", []):
            add_rich(rline(right, space_after=8), para)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    section_header(doc, "Education")
    for ed in cv["education"]:
        left, right = two_col(doc)
        left_label(left, ed["institution"], logo=ed.get("logo"))
        add_run(rline(right, first=True), f"{ed['degree']} - {ed['field']}", bold=True)
        add_run(rline(right, space_after=4),
                f"{fmt_date(ed['startDate'])} – {fmt_date(ed['endDate'])}", italic=True)
        if ed.get("gpa"):
            gp = rline(right, space_after=4)
            add_run(gp, "Grade: ", bold=True)
            add_run(gp, ed["gpa"])
        for ach in ed.get("achievements", []):
            add_run(rline(right, space_after=2), f"• {ach}")
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    section_header(doc, "Conferences and Events")
    for c in cv["conferences"]:
        left, right = two_col(doc)
        left_label(left, c["name"], logo=c.get("logo"))
        add_run(rline(right, first=True), c["type"], bold=True)
        add_run(rline(right, space_after=4), fmt_date(c["date"]), italic=True)
        if c.get("description"):
            add_run(rline(right, space_after=2), c["description"])
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(DOCX)
    print("DOCX written:", DOCX)
    to_pdf()


def to_pdf():
    try:
        import win32com.client as win32
    except ImportError:
        print("pywin32 not available; skipped PDF (DOCX is ready).")
        return
    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(DOCX))
        d.SaveAs(os.path.abspath(PDF), FileFormat=17)  # wdFormatPDF
        d.Close(False)
        print("PDF written:", PDF)
    except Exception as e:  # noqa: BLE001
        print("PDF step failed (DOCX still written):", e)
    finally:
        if word is not None:
            word.Quit()


if __name__ == "__main__":
    build()
